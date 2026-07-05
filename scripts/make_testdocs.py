"""Generate seeded-error test documents for the DiliAgent v1 review pipeline.

Produces three memo-style documents derived from `data/subset.json` eval
questions, each containing a mix of accurate claims (taken from gold answers),
corrupted claims (numbers/statements deliberately altered), and fabricated
claims (facts absent from the corpus). A `manifest.json` records the ground
truth per claim so the v1 annotation pipeline can be scored:

- `evals/testdocs/pepsico_memo.pdf`   (PDF — the annotation focus)
- `evals/testdocs/boeing_memo.docx`
- `evals/testdocs/amd_memo.md`
- `evals/testdocs/manifest.json`

Every `claim_text` appears verbatim in its document, so annotation anchoring
can be verified by exact string match.

Run:
    uv run --project backend --with python-docx scripts/make_testdocs.py
"""

from __future__ import annotations

import json
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
OUT_DIR = ROOT / "evals" / "testdocs"

# seeded_status -> expected verdict from the review pipeline
EXPECTED_VERDICT = {
    "accurate": "SUPPORTED",
    "corrupted": "CONTRADICTED",
    "fabricated": "NOT_IN_CORPUS",
}


def claim(
    claim_id: str,
    item_id: str | None,
    status: str,
    text: str,
    note: str = "",
) -> dict:
    return {
        "claim_id": claim_id,
        "item_id": item_id,
        "seeded_status": status,
        "expected_verdict": EXPECTED_VERDICT[status],
        "claim_text": text,
        "note": note,
    }


# ---------------------------------------------------------------------------
# Claim inventories. `claim_text` is embedded verbatim in the generated docs.
# ---------------------------------------------------------------------------

PEPSICO_SECTIONS: list[tuple[str, list[str | dict]]] = [
    (
        "Overview",
        [
            "PepsiCo is a leading global beverage and convenient food company "
            "with a portfolio spanning Pepsi-Cola, Frito-Lay, Gatorade, Quaker "
            "and SodaStream.",
            claim(
                "pep_c01",
                "pepsico_03",
                "accurate",
                "PepsiCo primarily operates across North America, Latin America, "
                "Europe, Africa, the Middle East, South Asia, Asia Pacific, "
                "Australia, New Zealand and China as of FY2022.",
            ),
            "The business is organized into divisions serving both developed "
            "and developing markets.",
        ],
    ),
    (
        "Credit and Liquidity",
        [
            "The company maintains committed credit facilities as backstop "
            "liquidity for its commercial paper program.",
            claim(
                "pep_c02",
                "pepsico_01",
                "corrupted",
                "On May 26, 2023, PepsiCo increased its unsecured five-year "
                "revolving credit agreement by $600 million.",
                "Gold: the increase was $400 million.",
            ),
            claim(
                "pep_c03",
                "pepsico_02",
                "accurate",
                "As of May 26, 2023, PepsiCo may borrow a total of $8.4 billion "
                "under its unsecured revolving credit agreements.",
            ),
        ],
    ),
    (
        "Earnings and Guidance",
        [
            claim(
                "pep_c04",
                "pepsico_04",
                "corrupted",
                "In Q1 FY2023, management raised full-year guidance for core "
                "constant currency EPS growth by 2 percentage points.",
                "Gold: guidance was raised by 1 percentage point.",
            ),
            claim(
                "pep_c05",
                "pepsico_07",
                "accurate",
                "The guidance raise reflected a strong start to FY2023.",
            ),
            claim(
                "pep_c06",
                "pepsico_05",
                "corrupted",
                "FY2022 unadjusted EBITDA less capital expenditures was "
                "approximately $10,342 million, defining unadjusted EBITDA as "
                "unadjusted operating income plus depreciation and amortization "
                "from the cash flow statement.",
                "Gold: $9,068 million.",
            ),
        ],
    ),
    (
        "Cost Structure",
        [
            "Management continues to execute multi-year productivity programs.",
            claim(
                "pep_c07",
                "pepsico_06",
                "accurate",
                "Restructuring costs outlined in the FY2022 income statement "
                "amounted to $411 million.",
            ),
        ],
    ),
    (
        "Governance and Workforce",
        [
            claim(
                "pep_c08",
                "pepsico_08",
                "corrupted",
                "At the AGM held on May 3, 2023, the shareholder proposal for a "
                "congruency report on net-zero emissions policies was approved "
                "by shareholders.",
                "Gold: the proposal was defeated.",
            ),
            claim(
                "pep_c09",
                None,
                "fabricated",
                "PepsiCo's global employee attrition rate improved to 6.1% in "
                "FY2022.",
                "Attrition rate is not disclosed in the corpus filings.",
            ),
        ],
    ),
    (
        "Capital Return",
        [
            claim(
                "pep_c10",
                None,
                "fabricated",
                "PepsiCo repurchased $12.5 billion of common shares during "
                "FY2024.",
                "The corpus contains no FY2024 filings; unverifiable by design.",
            ),
            "We view the dividend as well covered by operating cash flow.",
        ],
    ),
]

BOEING_SECTIONS: list[tuple[str, list[str | dict]]] = [
    (
        "Revenue Mix",
        [
            claim(
                "ba_c01",
                "boeing_01",
                "accurate",
                "Three categories each represented more than 20% of Boeing's "
                "FY2022 revenue: Commercial Airplanes at 39%, Defense at 35%, "
                "and Services at 26% of total revenue.",
            ),
            claim(
                "ba_c02",
                "boeing_07",
                "accurate",
                "Boeing's primary customers are a limited number of commercial "
                "airlines and the US government, with the US government "
                "accounting for 40% of total revenues in FY2022.",
            ),
        ],
    ),
    (
        "Margins and Tax",
        [
            claim(
                "ba_c03",
                "boeing_02",
                "corrupted",
                "Boeing's gross margin profile deteriorated in FY2022, with "
                "gross margin falling from 5.3% in FY2021 to 4.8% in FY2022.",
                "Gold: margin improved from 4.8% (FY2021) to 5.3% (FY2022).",
            ),
            claim(
                "ba_c04",
                "boeing_06",
                "accurate",
                "The effective tax rate was 0.62% in FY2022, compared with "
                "-14.76% in FY2021.",
            ),
        ],
    ),
    (
        "Balance Sheet",
        [
            claim(
                "ba_c05",
                "boeing_03",
                "corrupted",
                "Year-end FY2018 net property, plant, and equipment stood at "
                "$13,645 million.",
                "Gold: $12,645 million.",
            ),
        ],
    ),
    (
        "Outlook and Risk",
        [
            claim(
                "ba_c06",
                "boeing_04",
                "accurate",
                "Boeing's business remains subject to cyclicality driven by its "
                "exposure to the cyclical airline industry.",
            ),
            claim(
                "ba_c07",
                "boeing_05",
                "corrupted",
                "For 2023, Boeing is forecasting production rate cuts for the "
                "737 and 787 programs.",
                "Gold: production rate increases for the 737, 777X and 787.",
            ),
            claim(
                "ba_c08",
                None,
                "fabricated",
                "Boeing's win rate on competitive defense bids was approximately "
                "71% in FY2022.",
                "Bid win rate is not disclosed in the corpus filings.",
            ),
        ],
    ),
]

AMD_SECTIONS: list[tuple[str, list[str | dict]]] = [
    (
        "Liquidity",
        [
            claim(
                "amd_c01",
                "amd_01",
                "corrupted",
                "AMD's FY2022 quick ratio was 0.91, indicating a strained "
                "liquidity position.",
                "Gold: quick ratio 1.57, healthy liquidity.",
            ),
        ],
    ),
    (
        "Segments and Revenue Drivers",
        [
            claim(
                "amd_c02",
                "amd_02",
                "accurate",
                "Excluding Embedded, the Data Center segment showed the largest "
                "proportional sales increase from FY2021 to FY2022.",
            ),
            claim(
                "amd_c03",
                "amd_05",
                "accurate",
                "FY2022 revenue growth was driven by higher EPYC server "
                "processor sales, higher semi-custom product sales, and the "
                "inclusion of Xilinx embedded product sales.",
            ),
            claim(
                "amd_c04",
                None,
                "fabricated",
                "Average selling prices for EPYC processors rose approximately "
                "12% in FY2022.",
                "EPYC ASP change is not disclosed in the corpus filings.",
            ),
        ],
    ),
    (
        "Customers",
        [
            claim(
                "amd_c05",
                "amd_04",
                "accurate",
                "AMD reported customer concentration in FY2022, with one "
                "customer accounting for 16% of consolidated net revenue.",
            ),
        ],
    ),
    (
        "Historical Financials and Cash Flow",
        [
            claim(
                "amd_c06",
                "amd_03",
                "corrupted",
                "FY2015 depreciation and amortization margin, taking D&A from "
                "the cash flow statement as a percentage of revenue, was 7.5%.",
                "Gold: 4.2%.",
            ),
            claim(
                "amd_c07",
                "amd_06",
                "corrupted",
                "Financing activities brought in the most cash for AMD in "
                "FY2022, ahead of operating and investing activities.",
                "Gold: operating activities brought in the most cash.",
            ),
        ],
    ),
]

DOCS = [
    {
        "doc_id": "pepsico_memo",
        "format": "pdf",
        "filename": "pepsico_memo.pdf",
        "company": "PepsiCo",
        "title": "Project Cola — Draft Diligence Memo (PepsiCo, Inc.)",
        "intro": (
            "This draft memo summarizes preliminary findings on PepsiCo, Inc. "
            "ahead of the investment committee review. Figures are drawn from "
            "company filings and analyst workpapers and have not yet been "
            "independently verified."
        ),
        "sections": PEPSICO_SECTIONS,
    },
    {
        "doc_id": "boeing_memo",
        "format": "docx",
        "filename": "boeing_memo.docx",
        "company": "Boeing",
        "title": "Project Skyline — Draft Diligence Memo (The Boeing Company)",
        "intro": (
            "This draft memo summarizes preliminary findings on The Boeing "
            "Company. Figures are drawn from company filings and analyst "
            "workpapers and have not yet been independently verified."
        ),
        "sections": BOEING_SECTIONS,
    },
    {
        "doc_id": "amd_memo",
        "format": "markdown",
        "filename": "amd_memo.md",
        "company": "AMD",
        "title": "Project Silicon — Draft Diligence Memo (Advanced Micro Devices)",
        "intro": (
            "This draft memo summarizes preliminary findings on Advanced Micro "
            "Devices, Inc. Figures are drawn from company filings and analyst "
            "workpapers and have not yet been independently verified."
        ),
        "sections": AMD_SECTIONS,
    },
]


def section_paragraph(parts: list[str | dict]) -> str:
    """Join filler sentences and claim sentences into one flowing paragraph."""
    sentences = []
    for part in parts:
        sentences.append(part["claim_text"] if isinstance(part, dict) else part)
    return " ".join(sentences)


def render_markdown(doc: dict) -> str:
    lines = [f"# {doc['title']}", "", doc["intro"], ""]
    for heading, parts in doc["sections"]:
        lines.append(f"## {heading}")
        lines.append("")
        lines.append(section_paragraph(parts))
        lines.append("")
    return "\n".join(lines)


def render_pdf(doc: dict, path: Path) -> None:
    import fitz  # pymupdf

    body = [
        f"<h1>{doc['title']}</h1>",
        f"<p><i>{doc['intro']}</i></p>",
    ]
    for heading, parts in doc["sections"]:
        body.append(f"<h2>{heading}</h2>")
        body.append(f"<p>{section_paragraph(parts)}</p>")
    html = "".join(body)

    story = fitz.Story(html=html)
    writer = fitz.DocumentWriter(str(path))
    mediabox = fitz.paper_rect("letter")
    where = mediabox + (54, 54, -54, -54)
    more = 1
    while more:
        device = writer.begin_page(mediabox)
        more, _ = story.place(where)
        story.draw(device)
        writer.end_page()
    writer.close()


def render_docx(doc: dict, path: Path) -> None:
    import docx

    d = docx.Document()
    d.add_heading(doc["title"], level=0)
    d.add_paragraph(doc["intro"])
    for heading, parts in doc["sections"]:
        d.add_heading(heading, level=1)
        d.add_paragraph(section_paragraph(parts))
    d.save(str(path))


def build_manifest() -> dict:
    docs = []
    for doc in DOCS:
        claims = []
        for heading, parts in doc["sections"]:
            for part in parts:
                if isinstance(part, dict):
                    claims.append({**part, "section": heading})
        docs.append(
            {
                "doc_id": doc["doc_id"],
                "filename": doc["filename"],
                "format": doc["format"],
                "company": doc["company"],
                "title": doc["title"],
                "claims": claims,
            }
        )
    counts: dict[str, int] = {}
    for d in docs:
        for c in d["claims"]:
            counts[c["seeded_status"]] = counts.get(c["seeded_status"], 0) + 1
    return {
        "schema_version": 1,
        "description": (
            "Ground truth for the DiliAgent v1 seeded-error test documents. "
            "Each claim_text appears verbatim in its document. "
            "expected_verdict is what a correct review pipeline should report."
        ),
        "generator": "scripts/make_testdocs.py",
        "source_dataset": "data/subset.json",
        "totals": counts,
        "documents": docs,
    }


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    for doc in DOCS:
        path = OUT_DIR / doc["filename"]
        if doc["format"] == "pdf":
            render_pdf(doc, path)
        elif doc["format"] == "docx":
            render_docx(doc, path)
        else:
            path.write_text(render_markdown(doc), encoding="utf-8")
        print(f"wrote {path.relative_to(ROOT)}")

    manifest = build_manifest()
    manifest_path = OUT_DIR / "manifest.json"
    manifest_path.write_text(
        json.dumps(manifest, indent=2, ensure_ascii=False) + "\n", encoding="utf-8"
    )
    print(f"wrote {manifest_path.relative_to(ROOT)}  totals={manifest['totals']}")

    # Sanity: every claim must appear verbatim in its rendered document text.
    # NFKC folds the typographic ligatures (fi, fl) that PDF text extraction
    # yields — v1 annotation anchoring must be tolerant of these too.
    import unicodedata

    import fitz

    for doc, entry in zip(DOCS, manifest["documents"]):
        path = OUT_DIR / doc["filename"]
        if doc["format"] == "pdf":
            with fitz.open(str(path)) as pdf:
                text = " ".join(page.get_text() for page in pdf)
        elif doc["format"] == "docx":
            import docx

            text = " ".join(p.text for p in docx.Document(str(path)).paragraphs)
        else:
            text = path.read_text(encoding="utf-8")
        normalized = " ".join(unicodedata.normalize("NFKC", text).split())
        missing = [
            c["claim_id"]
            for c in entry["claims"]
            if " ".join(c["claim_text"].split()) not in normalized
        ]
        if missing:
            raise SystemExit(f"{doc['doc_id']}: claims not found verbatim: {missing}")
    print("verbatim check passed for all claims")


if __name__ == "__main__":
    main()
