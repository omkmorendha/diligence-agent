"""S6 DOCX + Markdown annotation tests (spec sections 1.8, 9).

Runs the two office annotators over the committed fixtures (boeing_memo.docx,
amd_memo.md) with synthetic verification results derived from manifest.json.

The mandatory guarantee (spec section 1.8): stripping the inserted `<mark>` tags
and the appended appendix reproduces the original Markdown file byte-for-byte.
"""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document

from app.review.annotate_docx import annotate_docx
from app.review.annotate_md import annotate_md, strip_annotations
from app.review.parse import anchor_quote, parse_document
from app.schemas import (
    Citation,
    Claim,
    ClaimValue,
    ReviewReport,
    ReviewReportClaim,
    ReviewSummary,
    VerificationResult,
)

TESTDOCS = Path(__file__).resolve().parents[2] / "evals" / "testdocs"
MANIFEST = json.loads((TESTDOCS / "manifest.json").read_text())

# Force a spread of verdicts across the color/branch map regardless of the
# manifest's own expected_verdict, so a single fixture exercises every path
# (green/red/amber/grey, the uncolored UNVERIFIABLE branch, and a result-less
# SKIPPED claim that must still land in the appendix).
_VERDICT_OVERRIDE = {
    1: "UNVERIFIABLE",
    2: "OUT_OF_SCOPE",
    3: "PARTIALLY_SUPPORTED",
}


def _doc_entry(filename: str) -> dict:
    return next(d for d in MANIFEST["documents"] if d["filename"] == filename)


def _citation(company: str) -> Citation:
    return Citation(
        citation_id="cit1",
        doc_id="doc1",
        doc_name=f"{company} 10-K",
        pdf_page=42,
        chunk_id="chunk1",
        quote="corpus evidence text",
        char_start=0,
        char_end=20,
    )


def _build_report(filename: str) -> tuple[ReviewReport, object]:
    """Parse a fixture and pair each manifest claim with a synthetic result."""
    entry = _doc_entry(filename)
    docmodel = parse_document(TESTDOCS / filename)
    report_claims: list[ReviewReportClaim] = []
    summary = ReviewSummary()
    for i, mc in enumerate(entry["claims"]):
        anchor = anchor_quote(docmodel, mc["claim_text"])
        assert anchor is not None, f"{mc['claim_id']} must anchor"
        verdict = _VERDICT_OVERRIDE.get(i, mc["expected_verdict"])
        claim = Claim(
            claim_id=mc["claim_id"],
            quote=anchor.quote,
            claim_type="numeric",
            company=entry["company"],
            question=f"Verify: {mc['claim_text'][:40]}",
            status="VERIFIED",
            anchor=anchor,
        )
        # Last claim is a result-less SKIPPED — appendix only, never marked.
        if i == len(entry["claims"]) - 1:
            claim.status = "SKIPPED"
            report_claims.append(ReviewReportClaim(claim=claim, result=None))
            summary.skipped += 1
            continue
        result = VerificationResult(
            claim_id=claim.claim_id,
            verdict=verdict,  # type: ignore[arg-type]
            explanation=f"Synthetic {verdict} explanation for {claim.claim_id}.",
            citations=[_citation(entry["company"])],
        )
        if verdict in ("CONTRADICTED", "PARTIALLY_SUPPORTED"):
            result.doc_value = ClaimValue(value=600.0, unit="USD millions")
            result.corpus_value = ClaimValue(value=400.0, unit="USD millions")
        elif verdict == "SUPPORTED":
            result.corpus_value = ClaimValue(value=8.4, unit="USD billions")
        elif verdict == "NOT_IN_CORPUS":
            result.queries_tried = ["q1", "q2", "q3"]
        report_claims.append(ReviewReportClaim(claim=claim, result=result))
    summary.total_claims = len(report_claims)
    report = ReviewReport(
        review_id="review_test_1",
        filename=filename,
        format=docmodel.format,
        company_scope=[entry["company"]],
        summary=summary,
        claims=report_claims,
    )
    return report, docmodel


# --- Markdown ---------------------------------------------------------------
def test_md_strip_reproduces_original_bytes(tmp_path: Path) -> None:
    """Spec section 1.8 mandatory: stripping inserted tags + appendix reproduces
    the original file byte-for-byte."""
    src = TESTDOCS / "amd_memo.md"
    original = src.read_bytes()
    report, _ = _build_report("amd_memo.md")
    out = tmp_path / "annotated.md"
    annotate_md(src, report, out)

    annotated = out.read_bytes().decode("utf-8")
    assert '<mark class="verdict-' in annotated  # something was actually inserted
    assert "## Review Appendix" in annotated

    restored = strip_annotations(annotated)
    assert restored.encode("utf-8") == original


def test_md_marks_carry_verdict_class_and_title(tmp_path: Path) -> None:
    src = TESTDOCS / "amd_memo.md"
    report, _ = _build_report("amd_memo.md")
    out = tmp_path / "annotated.md"
    annotate_md(src, report, out)
    annotated = out.read_text(encoding="utf-8")
    # A CONTRADICTED manifest claim produces a red-class mark with a hover title.
    assert 'class="verdict-contradicted"' in annotated
    assert 'title="' in annotated
    # Uncolored branch still emits a class (UNVERIFIABLE override on claim index 1).
    assert 'class="verdict-unverifiable"' in annotated
    # Appendix references the numbered marks.
    assert "[R1]" in annotated


def test_md_marks_are_well_nested(tmp_path: Path) -> None:
    src = TESTDOCS / "amd_memo.md"
    report, _ = _build_report("amd_memo.md")
    out = tmp_path / "annotated.md"
    annotate_md(src, report, out)
    body = out.read_text(encoding="utf-8").split("<!-- DILIAGENT")[0]
    assert body.count("<mark") == body.count("</mark>")


# --- DOCX -------------------------------------------------------------------
def _iter_runs(document: Document):
    for para in document.paragraphs:
        for run in para.runs:
            yield run


def test_docx_highlights_markers_and_appendix(tmp_path: Path) -> None:
    src = TESTDOCS / "boeing_memo.docx"
    report, _ = _build_report("boeing_memo.docx")
    original_para_count = len(Document(str(src)).paragraphs)
    out = tmp_path / "annotated.docx"
    annotate_docx(src, report, out)

    document = Document(str(out))
    # Appendix section was appended.
    assert len(document.paragraphs) > original_para_count
    assert any("Review Appendix" in p.text for p in document.paragraphs)

    # At least one run carries a highlight and a [R#] marker was inserted.
    highlighted = [r for r in _iter_runs(document) if r.font.highlight_color is not None]
    assert highlighted, "expected at least one highlighted run"
    full_text = "\n".join(p.text for p in document.paragraphs)
    assert "[R1]" in full_text


def test_docx_native_comments_added(tmp_path: Path) -> None:
    src = TESTDOCS / "boeing_memo.docx"
    report, _ = _build_report("boeing_memo.docx")
    out = tmp_path / "annotated.docx"
    annotate_docx(src, report, out)
    document = Document(str(out))
    if hasattr(document, "comments"):
        # python-docx >= 1.2: every marked claim gets a native comment.
        assert len(document.comments) >= 1


def test_docx_original_claim_text_preserved(tmp_path: Path) -> None:
    """Highlighting splits runs and inserts markers but never rewrites the source
    prose — every marked claim quote still anchors in the annotated document."""
    src = TESTDOCS / "boeing_memo.docx"
    report, _ = _build_report("boeing_memo.docx")
    out = tmp_path / "annotated.docx"
    annotate_docx(src, report, out)

    annotated_model = parse_document(out)
    for rc in report.claims:
        if rc.result is None:
            continue
        assert anchor_quote(annotated_model, rc.claim.quote) is not None, (
            f"{rc.claim.claim_id} quote lost after annotation"
        )
