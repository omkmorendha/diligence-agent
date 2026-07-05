"""S6 — DOCX annotation (spec sections 1.8, 9).

FROZEN CONTRACT — signature must not change.

Split runs at anchor boundaries, apply a `WD_COLOR_INDEX` highlight per verdict,
insert `[R1]`-style markers, and append a Review Appendix section. Use native
`add_comment` when python-docx >= 1.2; otherwise markers + appendix are the
guaranteed path. Writes the annotated copy to `out`.
"""

from __future__ import annotations

import copy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from ..schemas import ClaimValue, Citation, ReviewReport, ReviewReportClaim
from .parse import _match_quote_offsets

# Verdict -> highlight. UNVERIFIABLE stays uncolored (spec section 1.6) but still
# gets a marker + comment + appendix entry; OUT_OF_SCOPE / NOT_IN_CORPUS are grey.
_HIGHLIGHT: dict[str, WD_COLOR_INDEX | None] = {
    "SUPPORTED": WD_COLOR_INDEX.BRIGHT_GREEN,
    "CONTRADICTED": WD_COLOR_INDEX.RED,
    "PARTIALLY_SUPPORTED": WD_COLOR_INDEX.YELLOW,
    "NOT_IN_CORPUS": WD_COLOR_INDEX.GRAY_25,
    "OUT_OF_SCOPE": WD_COLOR_INDEX.GRAY_25,
    "UNVERIFIABLE": None,
}


def _fmt_value(v: ClaimValue | None) -> str:
    if v is None or v.value is None:
        return "n/a"
    num = f"{v.value:g}"
    return f"{num} {v.unit}" if v.unit else num


def _fmt_citation(citations: list[Citation]) -> str:
    if not citations:
        return ""
    c = citations[0]
    return f" (source: {c.doc_name} p.{c.pdf_page})"


def _comment_text(rc: ReviewReportClaim, ref: str) -> str:
    result = rc.result
    assert result is not None
    parts = [f"[{ref}] {result.verdict}"]
    if result.doc_value or result.corpus_value:
        parts.append(f"doc: {_fmt_value(result.doc_value)}; corpus: {_fmt_value(result.corpus_value)}")
    if result.explanation:
        parts.append(result.explanation)
    return " — ".join(parts) + _fmt_citation(result.citations)


# --- run splitting ----------------------------------------------------------
def _run_offsets(paragraph: Paragraph) -> list[tuple[Run, int, int]]:
    """(run, char_start, char_end) for each run over the paragraph's text."""
    offsets: list[tuple[Run, int, int]] = []
    pos = 0
    for run in paragraph.runs:
        offsets.append((run, pos, pos + len(run.text)))
        pos += len(run.text)
    return offsets


def _split_run_at(run: Run, k: int) -> None:
    """Split `run` at char offset k (0 < k < len(run.text)) into two adjacent
    runs, preserving formatting. The new right-hand run is inserted after it."""
    text = run.text
    new_r = copy.deepcopy(run._r)  # captures rPr before we mutate the original
    run.text = text[:k]
    new_run = Run(new_r, run._parent)
    new_run.text = text[k:]
    run._r.addnext(new_r)


def _ensure_boundary(paragraph: Paragraph, offset: int) -> None:
    """Guarantee a run boundary at char `offset` by splitting the straddling run."""
    for run, start, end in _run_offsets(paragraph):
        if start < offset < end:
            _split_run_at(run, offset - start)
            return


def _insert_marker_after(run: Run, text: str) -> None:
    """Insert an uncolored `[R#]` marker run immediately after `run`."""
    new_r = copy.deepcopy(run._r)
    new_run = Run(new_r, run._parent)
    new_run.text = text
    new_run.font.highlight_color = None
    run._r.addnext(new_r)


def _highlight_span(
    document: Document,
    paragraph: Paragraph,
    start: int,
    end: int,
    color: WD_COLOR_INDEX | None,
    marker: str,
    comment: str,
) -> bool:
    """Highlight [start, end) within `paragraph`, add a native comment when the
    library supports it, and append a marker. Returns True if anything covered."""
    _ensure_boundary(paragraph, end)
    _ensure_boundary(paragraph, start)
    covered = [run for run, s, e in _run_offsets(paragraph) if s >= start and e <= end and e > s]
    if not covered:
        return False
    if color is not None:
        for run in covered:
            run.font.highlight_color = color
    if comment and hasattr(document, "add_comment"):
        try:
            document.add_comment(covered, text=comment, author="DiliAgent", initials="DA")
        except Exception:
            # Native comments are best-effort (spec section 1.8); markers + appendix
            # are the guaranteed path, so a comment failure must not sink the review.
            pass
    _insert_marker_after(covered[-1], marker)
    return True


def _find_span(paragraphs: list[Paragraph], quote: str, hint: int | None) -> tuple[int, int, int] | None:
    """Locate `quote` within a single paragraph, preferring the hinted index.
    Returns (paragraph_index, start, end) or None."""
    order = []
    if hint is not None and 0 <= hint < len(paragraphs):
        order.append(hint)
    order.extend(i for i in range(len(paragraphs)) if i != hint)
    for i in order:
        match = _match_quote_offsets(paragraphs[i].text, quote)
        if match is not None:
            return i, match[0], match[1]
    return None


def _append_appendix(document: Document, report: ReviewReport, ref_by_id: dict[str, str]) -> None:
    document.add_page_break()
    document.add_heading("Review Appendix", level=1)
    document.add_paragraph(
        f"Automated review of {report.filename} — {report.summary.total_claims} claims."
    )
    for rc in report.claims:
        claim = rc.claim
        result = rc.result
        ref = ref_by_id.get(claim.claim_id) or claim.status
        verdict = result.verdict if result else claim.status
        para = document.add_paragraph()
        para.add_run(f"[{ref}] {verdict} — ").bold = True
        quote = claim.quote if len(claim.quote) <= 160 else claim.quote[:157] + "..."
        detail = f"“{quote}”"
        if result and (result.doc_value or result.corpus_value):
            detail += f" doc: {_fmt_value(result.doc_value)}; corpus: {_fmt_value(result.corpus_value)}."
        if result and result.explanation:
            detail += f" {result.explanation}"
        if result:
            detail += _fmt_citation(result.citations)
        para.add_run(detail)


def annotate_docx(src: str | Path, report: ReviewReport, out: str | Path) -> None:
    """Write an annotated copy of the source DOCX (highlights + markers) to `out`."""
    document = Document(str(src))
    paragraphs = list(document.paragraphs)

    ref_by_id: dict[str, str] = {}
    counter = 0
    for rc in report.claims:
        if rc.result is None or rc.claim.anchor is None:
            continue
        quote = rc.claim.anchor.quote or rc.claim.quote
        located = _find_span(paragraphs, quote, rc.claim.anchor.para_index)
        if located is None:
            continue
        para_index, start, end = located
        ref = f"R{counter + 1}"
        applied = _highlight_span(
            document,
            paragraphs[para_index],
            start,
            end,
            _HIGHLIGHT.get(rc.result.verdict),
            f" [{ref}]",
            _comment_text(rc, ref),
        )
        if applied:
            ref_by_id[rc.claim.claim_id] = ref
            counter += 1

    _append_appendix(document, report, ref_by_id)
    document.save(str(out))
